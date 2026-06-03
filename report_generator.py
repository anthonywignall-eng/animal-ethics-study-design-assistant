"""Build final report sections and generate a Word document."""

from __future__ import annotations

from datetime import date
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from ai_writer import build_rule_based_introduction
from config import ETHICS_DISCLAIMER
from flowchart import render_flowchart_png
from rules import evaluate_rules
from study_recommender import StudyRecommendation, recommend_study_design


NOT_SUPPLIED = "Not supplied"


def clean(value) -> str:
    if value is None:
        return NOT_SUPPLIED
    if isinstance(value, list):
        return ", ".join(value) if value else NOT_SUPPLIED
    text = str(value).strip()
    return text if text else NOT_SUPPLIED


def build_report_sections(response: dict) -> dict:
    evaluation = evaluate_rules(response)
    recommendation = recommend_study_design(response)
    project_introduction = response.get("ai_introduction") or build_rule_based_introduction(response, recommendation)
    sections = {
        "readiness": evaluation.readiness,
        "project_introduction": project_introduction,
        "recommendation": recommendation,
        "study_summary": build_study_summary(response),
        "test_articles": build_test_article_table(response),
        "animal_model": build_animal_model_table(response, evaluation),
        "route_schedule": build_route_schedule_table(response),
        "procedure_ethics": build_procedure_ethics_table(recommendation),
        "welfare": build_welfare_summary(response),
        "welfare_costs": recommendation.welfare_costs,
        "reduction": build_reduction_summary(response),
        "three_rs": recommendation.three_rs,
        "animal_timeline": recommendation.animal_timeline,
        "assumptions": recommendation.assumptions,
        "critical_missing": evaluation.all_missing,
        "warnings": evaluation.all_warnings_and_considerations,
        "draft_ethics": build_draft_ethics_sections(response, evaluation, recommendation),
    }
    return sections


def build_study_summary(response: dict) -> str:
    parts = []
    if response.get("study_title"):
        parts.append(f"The proposed study is titled '{response['study_title']}'.")
    if response.get("study_type"):
        parts.append(f"The proposed study type is {response['study_type'].lower()}.")
    if response.get("immediate_decision"):
        parts.append(f"The immediate decision to support is: {response['immediate_decision']}.")
    if response.get("downstream_study"):
        downstream = response.get("downstream_study_other") if response.get("downstream_study") == "Other" else response.get("downstream_study")
        parts.append(f"The planned downstream study is: {clean(downstream)}.")
    if response.get("test_article_types"):
        parts.append(f"The proposed test article categories are: {clean(response.get('test_article_types'))}.")
    if not parts:
        return "No study summary can be generated until core project information is supplied."
    return " ".join(parts)


def build_test_article_table(response: dict) -> list[dict]:
    rows = []
    details = response.get("test_article_details", {}) or {}
    for category in response.get("test_article_types", []):
        detail = details.get(category, {})
        rows.append(
            {
                "Test article": clean(detail.get("name") or category),
                "Category": category,
                "Purpose": clean(_article_purpose(detail)),
                "Formulation/vehicle status": clean(detail.get("formulation")),
                "Characterisation status": clean(detail.get("characterised")),
                "Dose-expression basis": clean(detail.get("adc_dose_expression")),
            }
        )
    return rows


def _article_purpose(detail: dict) -> str:
    return (
        detail.get("free_payload_purpose")
        or detail.get("prodrug_intended_use")
        or detail.get("adc_payload")
        or ""
    )


def build_animal_model_table(response: dict, evaluation=None) -> list[dict]:
    missing = []
    if response.get("sex") in {"Female only", "Male only"} and not response.get("sex_justification"):
        missing.append("Single-sex justification")
    if not response.get("strain"):
        missing.append("Strain")
    if not response.get("downstream_model") or "to be confirmed" in response.get("downstream_model", "").lower():
        missing.append("Downstream model")
    return [
        {
            "Species": clean(response.get("species")),
            "Strain": clean(response.get("strain")),
            "Sex": clean(response.get("sex")),
            "Healthy or tumour-bearing": clean(response.get("animal_status")),
            "Relevance to downstream study": clean(response.get("model_match")),
            "Missing justification items": clean(missing),
        }
    ]


def build_route_schedule_table(response: dict) -> list[dict]:
    return [
        {
            "Route": clean(response.get("route")),
            "MTD schedule": clean(response.get("mtd_schedule_other") if response.get("mtd_schedule") == "Other" else response.get("mtd_schedule")),
            "Intended downstream schedule": clean(response.get("downstream_schedule_other") if response.get("downstream_schedule") == "Other" else response.get("downstream_schedule")),
            "Whether schedules align": clean(response.get("schedule_alignment")),
            "Prior evidence supplied": clean(response.get("prior_tolerability")),
        }
    ]


def build_procedure_ethics_table(recommendation: StudyRecommendation) -> list[dict]:
    rows = []
    for item in recommendation.procedure_items:
        rows.append(
            {
                "Procedure": item.procedure,
                "Why included": item.why_included,
                "Associated pain/distress": item.associated_pain_or_distress,
                "Anaesthesia/analgesia consideration": item.anaesthesia_or_analgesia_consideration,
                "Welfare cost": item.welfare_cost,
            }
        )
    return rows


def build_welfare_summary(response: dict) -> dict:
    monitoring_frequency = {
        "Routine monitoring": clean(response.get("monitoring_routine")),
        "Immediately after dosing": clean(response.get("monitoring_immediate")),
        "Expected toxicity window": clean(response.get("monitoring_toxicity_window")),
        "After adverse signs": clean(response.get("monitoring_after_signs")),
    }
    return {
        "Anticipated adverse effects": clean(response.get("anticipated_adverse_effects")),
        "Monitoring parameters": clean(response.get("monitoring_parameters")),
        "Monitoring frequency": monitoring_frequency,
        "Humane endpoints": clean(response.get("clinical_scoring")),
        "Dose stopping rules": clean(response.get("cohort_toxicity_action")),
        "Supportive care": clean(response.get("supportive_care")),
        "Authorised decision-maker": clean(response.get("decision_maker")),
    }


def build_reduction_summary(response: dict) -> dict:
    return {
        "Staged escalation intended": clean(response.get("study_design")),
        "Lead-candidate prioritisation possible": clean(response.get("candidate_prioritisation")),
        "Later arms conditional": clean(response.get("conditional_testing")),
        "Controls requiring justification": clean(response.get("vehicle_control_justification")),
        "Animal number minimisation": clean(response.get("animal_number_minimisation")),
    }


def build_draft_ethics_sections(response: dict, evaluation, recommendation: StudyRecommendation) -> dict:
    missing_prompt = "Investigator confirmation is required for: " + clean(evaluation.all_missing)
    procedures_text = " ".join(
        f"{item.procedure}: {item.associated_pain_or_distress} {item.anaesthesia_or_analgesia_consideration}"
        for item in recommendation.procedure_items
    )
    welfare_cost_text = " ".join(recommendation.welfare_costs)
    timeline_text = " ".join(recommendation.animal_timeline)
    return {
        "Scientific aim": _paragraph(
            [
                f"Draft prompt: This study aims to support the following decision: {response.get('immediate_decision')}.",
                f"The proposed downstream study is: {response.get('downstream_study')}.",
                f"The recommended design pathway is: {recommendation.pathway}.",
            ]
        ),
        "Justification for animal use": clean(response.get("replacement_rationale"))
        if response.get("replacement_rationale")
        else "Draft prompt: Explain why in vivo tolerability information is required before the planned downstream animal study, and identify the non-animal information already used for candidate or dose-strategy selection.",
        "Justification for MTD/tolerability design": _paragraph(
            [
                f"Draft prompt: The intended outcome is {response.get('intended_outcome')}.",
                f"The working tolerability definition is: {response.get('tolerability_definition')}.",
                f"Proposed dose levels or range supplied by the investigator: {response.get('dose_levels')}.",
                f"Starting-dose rationale supplied: {response.get('starting_dose_evidence_text') or response.get('starting_dose_evidence')}.",
            ]
        ),
        "Proposed experimental approach": _paragraph(
            [
                f"Draft prompt: Test article category or categories: {clean(response.get('test_article_types'))}.",
                f"Animal model: {clean(response.get('species'))}, strain {clean(response.get('strain'))}, sex {clean(response.get('sex'))}.",
                f"Route: {clean(response.get('route'))}. MTD schedule: {clean(response.get('mtd_schedule'))}.",
                f"Study pathway: {recommendation.pathway}.",
            ]
        ),
        "Procedures and associated pain/distress": procedures_text,
        "Welfare cost to animals": welfare_cost_text,
        "Animal timeline from arrival to study completion": timeline_text,
        "Anticipated adverse effects and animal welfare impact": clean(response.get("anticipated_adverse_effects")),
        "Monitoring and humane endpoints": _paragraph(
            [
                f"Monitoring parameters: {clean(response.get('monitoring_parameters'))}.",
                f"Clinical scoring/intervention criteria: {clean(response.get('clinical_scoring'))}.",
                f"Individual endpoint action: {clean(response.get('individual_endpoint_action'))}.",
                f"Cohort stopping rule: {clean(response.get('cohort_toxicity_action'))}.",
            ]
        ),
        "Replacement": recommendation.three_rs.get("Replacement", clean(response.get("replacement_rationale"))),
        "Reduction": recommendation.three_rs.get("Reduction", clean(response.get("reduction_measures") or response.get("animal_number_minimisation"))),
        "Refinement": recommendation.three_rs.get("Refinement", clean(response.get("refinement_measures"))),
        "Information still requiring investigator confirmation": missing_prompt,
    }


def _paragraph(parts: list[str]) -> str:
    filtered = [part for part in parts if part and "None" not in part]
    return " ".join(filtered) if filtered else NOT_SUPPLIED


def generate_docx_report(response: dict) -> BytesIO:
    sections = build_report_sections(response)
    recommendation = sections["recommendation"]
    document = Document()
    _configure_document(document)

    title = document.add_heading("Animal Ethics Study Design Assistant", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph("MTD Module structured study-design report")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(f"Date generated: {date.today().isoformat()}")
    document.add_paragraph(ETHICS_DISCLAIMER).runs[0].bold = True
    document.add_page_break()

    document.add_heading("A. Readiness Classification", level=1)
    document.add_paragraph(sections["readiness"])

    document.add_heading("B. Project Introduction and Aims", level=1)
    document.add_paragraph(clean(sections["project_introduction"]))

    document.add_heading("C. Recommended Study Design", level=1)
    document.add_paragraph(recommendation.summary)
    document.add_paragraph("Recommendation rationale:").runs[0].bold = True
    _add_bullets(document, recommendation.rationale or ["No recommendation rationale supplied."])
    document.add_paragraph("Rule-based study flowchart:").runs[0].bold = True
    flowchart_image = render_flowchart_png(recommendation.flowchart_steps)
    document.add_picture(flowchart_image, width=Inches(6.5))

    document.add_heading("D. Proposed Study Summary", level=1)
    document.add_paragraph(sections["study_summary"])

    document.add_heading("E. Proposed Test Articles", level=1)
    _add_table(document, sections["test_articles"])

    document.add_heading("F. Proposed Animal Model", level=1)
    _add_table(document, sections["animal_model"])

    document.add_heading("G. Route and Schedule", level=1)
    _add_table(document, sections["route_schedule"])

    document.add_heading("H. Animal Ethics Procedure Considerations", level=1)
    document.add_paragraph(
        "The following procedure table is rule-based and should be checked against the final protocol, facility SOPs, veterinary advice and AEC requirements."
    )
    _add_table(document, sections["procedure_ethics"])

    document.add_heading("I. Welfare Cost to Animals", level=1)
    _add_bullets(document, sections["welfare_costs"])

    document.add_heading("J. Animal Timeline", level=1)
    _add_numbered_list(document, sections["animal_timeline"])

    document.add_heading("K. Monitoring and Welfare Framework", level=1)
    _add_key_value_section(document, sections["welfare"])

    document.add_heading("L. Proposed Animal-use Reduction Strategy", level=1)
    _add_key_value_section(document, sections["reduction"])

    document.add_heading("M. Replacement, Reduction and Refinement", level=1)
    _add_key_value_section(document, sections["three_rs"])

    document.add_heading("N. Critical Missing Information", level=1)
    _add_bullets(document, sections["critical_missing"] or ["No critical missing information identified by the current rule set."])

    document.add_heading("O. Design Warnings and Considerations", level=1)
    _add_bullets(document, sections["warnings"] or ["No design warnings or considerations identified by the current rule set."])

    document.add_heading("P. Assumptions and Protocol Items Requiring Confirmation", level=1)
    _add_bullets(document, sections["assumptions"])

    document.add_heading("Q. Draft Ethics-application Content", level=1)
    for heading, text in sections["draft_ethics"].items():
        document.add_heading(heading, level=2)
        document.add_paragraph(clean(text))

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)


def _add_table(document: Document, rows: list[dict]) -> None:
    if not rows:
        document.add_paragraph(NOT_SUPPLIED)
        return
    headers = list(rows[0].keys())
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, header in enumerate(headers):
            cells[index].text = clean(row.get(header))


def _add_key_value_section(document: Document, values: dict) -> None:
    for key, value in values.items():
        if isinstance(value, dict):
            document.add_paragraph(key).runs[0].bold = True
            for nested_key, nested_value in value.items():
                document.add_paragraph(f"{nested_key}: {clean(nested_value)}", style=None)
        else:
            document.add_paragraph(f"{key}: {clean(value)}")


def _add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(clean(item), style="List Bullet")


def _add_numbered_list(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(clean(item), style="List Number")
